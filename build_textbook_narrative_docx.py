from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name(
    "Basic_Circuits_and_Raspberry_Pi_Configuration_Textbook_Narrative.docx"
)


UNITS = [
    (
        "Unit 01",
        "Raspberry Pi Configuration",
        [
            ("Raspberry Pi in IoT", "Raspberry Pi is a small single-board computer that brings together computing, networking, and hardware control. In an IoT laboratory, it is valuable because the learner can use a Linux operating system, write Python programs, connect to the internet, and interact with real electronic circuits through GPIO pins. This makes the board a bridge between software and the physical world."),
            ("Raspberry Pi Board Overview", "A Raspberry Pi board includes a processor, memory, USB ports, video output, microSD storage, networking capability, and a GPIO header. The microSD card stores the operating system and user files. USB ports support keyboard, mouse, and other peripherals. Network interfaces allow the board to communicate with other devices and online services."),
            ("GPIO Header and Pin Groups", "The GPIO header is the most important hardware interface for basic IoT practice. It includes digital input/output pins, power supply pins, ground pins, and special-purpose pins used for communication. The course material identifies all-round I/O with digital pins, 5 V and 3.3 V supply pins, ground pins, and ID EEPROM pins. Correct pin identification is essential because a wrong connection can damage the board."),
            ("Preparing the SD Card", "Raspberry Pi usually boots from a microSD card. Preparing the card means selecting the correct operating system image, writing it with an imaging tool, and safely ejecting it after the write process is complete. The card should not be removed while writing is in progress. A corrupted SD card can prevent the board from booting."),
            ("Installing Raspberry Pi OS", "The operating system provides the desktop environment, terminal, drivers, package manager, and programming tools. A successful installation gives the learner a stable base for all later experiments. During installation, the user should choose the correct board model, storage device, and configuration settings."),
            ("First Boot and Basic Setup", "On first boot, Raspberry Pi may ask for region, keyboard, password, network, and update settings. These details are not minor formalities. A wrong keyboard setting can make passwords difficult to type, and an unfinished network setup can prevent package installation or remote access."),
            ("Raspberry Pi Configuration Tool", "The Raspberry Pi Configuration tool is used to enable system interfaces and adjust preferences. In the course flow, SSH and VNC are enabled from the configuration interface. These settings allow the board to be controlled remotely from another computer."),
            ("SSH and IP Address", "SSH provides remote command-line access. To use it, the Raspberry Pi must be connected to the same reachable network, SSH must be enabled, and the user must know the board's IP address. The course notes mention using ifconfig and checking the inet address under wlan0. If the device disconnects and reconnects, the IP address may change, so it should be checked again."),
            ("Using PuTTY for SSH", "PuTTY is a common Windows SSH client. The user enters the Raspberry Pi IP address, opens a session, accepts the host key when appropriate, and logs in with the username and password. After login, the terminal can be used to list files, change folders, install packages, and run Python scripts."),
            ("VNC Remote Desktop", "VNC provides remote desktop access. Unlike SSH, which gives a text terminal, VNC displays the graphical Raspberry Pi desktop on another computer. This is helpful when the learner wants to use graphical menus, editors, and file-transfer tools without connecting a separate monitor to the board."),
            ("File Transfer with VNC", "The course material shows creating a folder on the Windows PC, preparing a test text file, and transferring it to Raspberry Pi through VNC. This workflow is useful for moving scripts, notes, and experiment files. Clear folder names make file management easier during laboratory sessions."),
            ("Linux Terminal Basics", "The Linux terminal is central to Raspberry Pi work. Commands such as cd, ls, sudo apt-get update, and python3 are used repeatedly. Command spelling, spacing, and directory location matter. A command typed in the wrong folder may fail even when the software is installed correctly."),
            ("Software Installation and Safety", "Before installing packages, the package list should be updated so the system can find current software information. At the same time, hardware safety must be remembered. The Raspberry Pi should be shut down properly before power removal, and GPIO pins should never be connected to unsafe voltages."),
            ("Troubleshooting Remote Access", "Remote access problems usually come from disabled SSH or VNC, wrong IP address, network mismatch, incorrect password, or firewall restrictions. A disciplined troubleshooting method checks power, network, interface settings, address, username, and password before assuming the board has failed."),
        ],
    ),
    (
        "Unit 02",
        "Breadboard and Basic Sensor Circuits",
        [
            ("Introduction to Breadboards", "A breadboard is a reusable prototyping board used to build circuits without soldering. Components and jumper wires are inserted into holes that are internally connected in groups. Breadboards are important for beginners because they allow quick testing, correction, and part reuse."),
            ("Internal Connections", "A breadboard has hidden metal strips below the holes. Terminal strips are usually connected in short horizontal groups, while power rails run along the sides. The central groove separates the two halves and is useful for placing integrated circuits so that opposite pins are not shorted together."),
            ("Power and Ground Rails", "Power rails distribute supply voltage and ground across the board. Many breadboards split these rails into sections, so a rail that appears continuous may not actually be connected from end to end. Before testing a circuit, the learner should know exactly where voltage and ground are available."),
            ("Component Placement", "Good placement makes a circuit easier to understand and safer to debug. Components should be inserted so that their leads do not accidentally share the same connected row unless that connection is intended. Integrated circuits usually straddle the central groove."),
            ("Jumper Wires", "Jumper wires create electrical paths between breadboard points, Raspberry Pi pins, and components. Color discipline improves readability. Red is commonly used for positive supply, black or blue for ground, and other colors for signals. Short, neat wires reduce confusion."),
            ("Breadboard Learning Value", "The source course contrasts breadboard practice with soldered wiring. Without a breadboard, wires often need to be soldered together. With a breadboard, components can be reused and rearranged. This encourages experimentation while reducing cost and permanent mistakes."),
            ("Resistors on Breadboards", "Resistors limit current, divide voltage, and create defined input states. In LED circuits, a resistor prevents excessive current. In switch and sensor circuits, resistors may pull a signal line up to supply voltage or down to ground."),
            ("LED Circuit Basics", "An LED is polarity-sensitive. Current flows from anode to cathode when connected correctly. A current-limiting resistor is normally required. If the LED does not glow, the cause may be reversed polarity, wrong resistor placement, missing ground, or no output signal."),
            ("Pull-Up and Pull-Down Resistors", "A pull-up resistor connects a signal line to a positive supply so that the default input state is HIGH. A pull-down resistor connects a signal line to ground so that the default state is LOW. These resistors prevent floating inputs, which can behave unpredictably."),
            ("DHT Sensor Family", "DHT11 and DHT22/AM2302 sensors measure temperature and humidity. They are common in introductory IoT work because they combine sensing and digital data output in a small module. The DHT11 is simpler and lower-cost, while the DHT22 generally offers better range and accuracy."),
            ("DHT Pin Functions", "DHT sensors commonly use VCC, GND, DATA, and sometimes an unused pin. The course slide text shows VCC, DATA, not used, and ground labels. The exact pin order can vary by package or breakout module, so the module marking or datasheet should always be checked."),
            ("DHT Pull-Up Requirement", "The DHT data line often requires a pull-up resistor, commonly in the range of 5 k ohm to 10 k ohm. Some modules already include this resistor. If the pull-up is missing, the data signal may be unstable or unreadable."),
            ("Building a DHT Circuit", "A basic DHT circuit connects VCC to the correct supply voltage, GND to ground, and DATA to a GPIO pin. When needed, the pull-up resistor connects DATA to VCC. The ground of the sensor and Raspberry Pi must be common for the signal to be interpreted correctly."),
            ("Breadboard Debugging", "Debugging a breadboard circuit means checking physical reality before changing software. The learner should inspect row positions, rail connections, wire continuity, component polarity, and pin numbering. Many circuit failures are caused by one hole of misplacement."),
            ("Clean Wiring Practice", "Clean wiring is not only for appearance. It reduces mistakes, makes circuits easier to explain, and helps others review the work. A good circuit layout keeps related components near each other and avoids tangled wires crossing every part of the board."),
        ],
    ),
    (
        "Unit 03",
        "Sensor Fundamentals",
        [
            ("Meaning of a Sensor", "A sensor detects a physical quantity and converts it into an electrical signal or digital value. Temperature, humidity, distance, light, motion, pressure, and vibration can all be sensed. In IoT systems, sensors are the entry point for real-world data."),
            ("Input and Output", "The input of a sensor is the physical quantity being measured. The output may be voltage, current, resistance, frequency, pulse width, or digital data. Understanding the output type is necessary before choosing a Raspberry Pi interface."),
            ("Analog Sensors", "Analog sensors produce continuously varying signals. A temperature sensor may change voltage gradually as temperature changes, and a potentiometer may output a variable voltage. Since Raspberry Pi GPIO pins are digital, analog readings require an analog-to-digital converter."),
            ("Digital Sensors", "Digital sensors provide logic-level signals or structured communication. A digital motion sensor may output HIGH when motion is detected. Other sensors send measurements through protocols such as I2C, SPI, UART, or timed one-wire style communication."),
            ("Range", "Range describes the minimum and maximum value a sensor can measure. A sensor should be selected so that the expected measurement stays within its useful range. If the physical value goes beyond the range, the output may saturate or become unreliable."),
            ("Sensitivity", "Sensitivity describes how much the output changes for a given change in input. High sensitivity is useful for small changes, but it can also make the system more affected by noise. Sensitivity must be considered together with stability and resolution."),
            ("Accuracy and Error", "Accuracy describes how close a reading is to the true value. Error may come from poor calibration, sensor limitations, environmental changes, wiring problems, or electrical noise. A measured value should always be interpreted with possible error in mind."),
            ("Linearity", "A linear sensor changes output in direct proportion to input. Nonlinear sensors require formulas, lookup tables, or calibration curves. The learner should not assume that every sensor output changes in a perfectly straight-line relationship."),
            ("Hysteresis", "Hysteresis occurs when the sensor output for a value depends on whether the input is increasing or decreasing. The course material includes a hysteresis expression using full-scale output. This idea matters when repeated measurements do not return to exactly the same reading."),
            ("Response Time", "Response time is the delay between a change in input and the sensor output reaching the expected value. The course material shows t0.9, which represents the time required to reach about 90 percent of the final response. Slow response can be a problem in fast-changing systems."),
            ("Calibration", "Calibration relates sensor output to known reference values. For example, a temperature sensor can be compared with a trusted thermometer. Calibration does not make a poor sensor perfect, but it improves the reliability of interpretation."),
            ("Noise and Stability", "Noise is unwanted variation in readings. Stability is the ability to produce consistent readings under the same conditions. Good wiring, filtering, averaging, shielding, and stable power can improve measurement quality."),
            ("Sensor Selection", "A sensor should be selected by considering range, accuracy, sensitivity, response time, output type, supply voltage, cost, physical size, and environmental conditions. The best sensor is not always the most expensive one; it is the one that fits the application."),
            ("Temperature and Humidity Sensors", "Temperature sensors may use resistance changes, semiconductor behavior, or digital conversion. Humidity sensors measure moisture in air. DHT modules combine temperature and humidity sensing, making them useful in beginner IoT laboratories."),
            ("Motion and Distance Sensors", "Motion sensors detect movement or presence, while distance sensors estimate separation between the sensor and an object. These sensors are common in automation, security, robotics, and interactive projects."),
        ],
    ),
    (
        "Unit 04",
        "Basic Electronics and Circuits",
        [
            ("Water Analogy for Electricity", "The course uses water current and electric current as an analogy. Voltage can be imagined as pressure, current as flow, and resistance as restriction. The analogy is useful for beginners, but electrical circuits also have behaviors that water systems do not fully represent."),
            ("Voltage and Current", "Voltage is electric potential difference. It is measured between two points. Current is the flow of electric charge through a path. A circuit must be closed for steady current to flow. Both voltage and current must remain within component ratings."),
            ("Resistance and Ohm's Law", "Resistance opposes current. Ohm's law states that voltage equals current multiplied by resistance: V = I x R. This relationship helps calculate resistor values, current through LEDs, and expected voltage drops in simple circuits."),
            ("DC and AC", "Direct current maintains one polarity with time, while alternating current changes direction periodically. Batteries and Raspberry Pi supplies provide DC. The course material shows voltage versus time for DC and AC and introduces frequency as repeated vibration or cycles per second."),
            ("Frequency and Hertz", "Frequency is measured in hertz. One cycle per second is 1 Hz, two cycles per second is 2 Hz, and one thousand cycles per second is 1 kHz. Frequency matters in signals, communication, filtering, and timing circuits."),
            ("Series and Parallel Circuits", "In a series circuit, components are connected one after another and carry the same current. In a parallel circuit, components share the same two nodes and have the same voltage across them. Equivalent resistance differs in each arrangement."),
            ("Dry Cells and Supply Voltage", "Dry cells provide a fixed voltage such as 1.5 V. Cells can be combined to obtain higher voltage when connected properly. Polarity must be respected. Reversed supply connections can damage components or prevent circuits from operating."),
            ("Capacitors", "A capacitor stores electric charge and resists sudden changes in voltage. Capacitors appear in smoothing circuits, timer circuits, waveform shaping circuits, time-lag circuits, low-pass filters, and high-pass filters, as indicated in the course material."),
            ("Smoothing and Timing", "A smoothing capacitor reduces rapid voltage variation. In timer circuits, a capacitor charges and discharges through a resistor, creating a delay. This resistor-capacitor behavior is a foundation for many simple analog timing circuits."),
            ("Low-Pass and High-Pass Filters", "A low-pass filter allows slower signal changes to pass while reducing high-frequency noise. A high-pass filter allows faster changes to pass while reducing slow or DC components. Filter choice depends on the signal that should be preserved."),
            ("Diodes and Diode Types", "A diode has an anode and cathode and conducts more easily in one direction. The course lists general-purpose germanium diodes, silicon diodes, Schottky barrier diodes, and Zener diodes. Each type has a different practical role."),
            ("Zener Diodes", "A Zener diode is commonly used for voltage reference or voltage limiting when operated in its reverse breakdown region within rating. It is useful when a circuit needs a controlled voltage threshold."),
            ("Transistors", "A transistor can act as an amplifier or electronic switch. Bipolar transistors include NPN and PNP types, with collector, base, and emitter terminals. Field-effect transistors use a gate-controlled channel and are common in switching applications."),
            ("Back EMF Protection", "Coils and relay loads can generate reverse voltage when switched off. The course material shows adding a diode to absorb this reverse electromotive force. This protection diode prevents damage to the driving digital circuit."),
            ("Circuit Symbols and Connections", "Circuit diagrams use standardized symbols. The course notes explain that a semicircle can mean crossed wires, a dot means connected, and a T-shape means connected. Units such as ohms or microfarads may sometimes be omitted in diagrams."),
            ("Operational Amplifiers", "An operational amplifier has differential inputs and an output. It uses supply connections and can be configured as an amplifier, comparator, or signal conditioner. Feedback resistors determine much of its circuit behavior."),
            ("Inverting and Non-Inverting Amplifiers", "In an inverting amplifier, the output changes opposite to the input phase. In a non-inverting amplifier, the output follows the input phase while being amplified. Both use feedback to set gain."),
            ("Comparators", "A comparator compares two input voltages and changes its output depending on which input is higher. Comparators are useful for threshold decisions, such as converting an analog sensor level into a digital decision."),
            ("Digital ICs", "Digital integrated circuits operate using logic levels. They include gates, flip-flops, decoders, counters, buffers, and interface devices. Digital ICs are building blocks for larger electronic systems."),
            ("Flip-Flops", "A flip-flop stores one bit of information. A D flip-flop transfers input data to output under clock control. Flip-flops are used in memory, counters, registers, and timing logic."),
            ("Address and Data Buses", "A bus is a group of signal lines that work together. An address bus selects a device or location, while a data bus carries information. The course material shows address bus, data bus, IORQ, RD, WR, decoders, and buffers."),
            ("Decoders and Bus Buffers", "A decoder such as 74HC138 selects one output from several address inputs. A bus buffer such as 74HC245 or 74HC541 helps isolate or drive bus signals. These devices are common in microprocessor interface circuits."),
        ],
    ),
    (
        "Unit 05",
        "GPIO Zero Programming",
        [
            ("Introduction to GPIO Zero", "GPIO Zero is a Python library designed to simplify Raspberry Pi physical computing. It provides classes for common devices such as LEDs, buttons, motion sensors, and distance sensors, allowing beginners to focus on behavior rather than low-level pin control."),
            ("Installing GPIO Zero", "The course material shows updating package information and installing python3-gpiozero. Installation prepares the Raspberry Pi to run the example programs. If the package is missing, Python will report an import error."),
            ("Blinking One LED", "The first LED program turns an LED on and off with a delay. This simple experiment teaches output control, timing, and the relationship between a Python object and a GPIO pin."),
            ("LED Program with Functions", "Functions make code reusable. Instead of repeating the same instructions, a function can describe a behavior such as blinking. This is the first step toward readable and maintainable hardware programs."),
            ("Blinking Three LEDs", "Controlling three LEDs introduces pin organization and sequencing. The learner must keep track of which LED is connected to which GPIO pin and how timing changes the visible pattern."),
            ("MCP3008 and Potentiometer", "Raspberry Pi GPIO pins do not directly read analog voltage. The MCP3008 analog-to-digital converter allows analog values, such as those from a potentiometer, to be converted into digital readings that Python can process."),
            ("Button Input and Button Game", "A button is a digital input. GPIO Zero can detect pressed and released states. A button game combines input detection with timing and program logic, turning a simple switch into an interactive exercise."),
            ("Motion Sensor Program", "A motion sensor such as a PIR sensor outputs a signal when movement is detected. A GPIO Zero program can respond by printing a message, turning on an LED, or triggering another action."),
            ("Distance Sensor Program", "A distance sensor measures how far an object is from the sensor. Ultrasonic distance sensors use trigger and echo pulses. GPIO Zero hides much of the timing detail and provides distance readings in a beginner-friendly form."),
            ("Good Python Hardware Style", "Good hardware programs use clear variable names, fixed pin definitions, comments where helpful, and safe shutdown behavior. The program should be easy to match with the circuit on the breadboard."),
        ],
    ),
]


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 25, "0B2545", 0, 8),
        ("Heading 1", 18, "1F4D78", 18, 8),
        ("Heading 2", 14, "2E74B5", 12, 5),
        ("Heading 3", 12, "1F4D78", 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Textbook Note" not in styles:
        style = styles.add_style("Textbook Note", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor.from_string("444444")
        style.paragraph_format.left_indent = Inches(0.22)
        style.paragraph_format.right_indent = Inches(0.15)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(7)

    header = section.header.paragraphs[0]
    header.text = "Basic Circuits and Raspberry Pi Configuration for IoT Applications"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string("666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Page ").font.size = Pt(8.5)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def p(doc, text, style="Normal"):
    para = doc.add_paragraph(style=style)
    para.add_run(text)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        para.add_run(item)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc, rows):
    tbl = doc.add_table(rows=1, cols=len(rows[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, text in enumerate(rows[0]):
        cell = tbl.rows[0].cells[i]
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        shade(cell, "F2F4F7")
    for row in rows[1:]:
        cells = tbl.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p(doc, "", "Textbook Note")


def front_matter(doc):
    p(doc, "Basic Circuits and Raspberry Pi Configuration for IoT Applications", "Title")
    p(doc, "A textbook based on SIC IoT Chapter 2: Basic Circuit and Raspberry Pi Configuration", "Heading 2")
    p(
        doc,
        "This book introduces the hardware foundations needed for beginner Internet of Things practice. "
        "It follows the course path from Raspberry Pi configuration to breadboard circuits, sensors, electronics, and Python-based GPIO programming.",
    )
    p(
        doc,
        "The material is written as an original textbook manuscript. It does not reproduce the slide material word-for-word. "
        "Instead, it explains the same units and subunits in continuous textbook style, with clear headings, explanatory paragraphs, and selected reference tables.",
    )
    table(
        doc,
        [
            ["Part", "Coverage"],
            ["Unit 01", "Raspberry Pi setup, configuration, SSH, VNC, IP address, PuTTY, terminal use, and file transfer"],
            ["Unit 02", "Breadboard structure, power rails, component placement, resistors, LEDs, pull-up/pull-down circuits, and DHT sensors"],
            ["Unit 03", "Sensor meaning, analog and digital outputs, range, sensitivity, accuracy, linearity, hysteresis, response time, and calibration"],
            ["Unit 04", "Voltage, current, resistance, Ohm's law, AC/DC, frequency, capacitors, diodes, transistors, op-amps, digital ICs, buses, and buffers"],
            ["Unit 05", "GPIO Zero installation and programs for LEDs, buttons, MCP3008, potentiometer, motion sensor, and distance sensor"],
        ],
    )
    doc.add_page_break()

    p(doc, "Preface", "Heading 1")
    p(
        doc,
        "IoT education begins with the ability to connect a computing board to the physical world. A student must learn not only programming, "
        "but also voltage, current, signal pins, component placement, sensor behavior, and practical troubleshooting. Raspberry Pi is a strong teaching platform because it combines a familiar computer environment with accessible hardware control.",
    )
    p(
        doc,
        "This textbook is arranged in five units. The first unit prepares the Raspberry Pi. The second develops breadboard and sensor-circuit skill. "
        "The third explains how sensors behave and how their performance is judged. The fourth provides the electronics background needed to read and build simple circuits. The final unit uses GPIO Zero to control and read hardware through Python.",
    )
    p(
        doc,
        "The tone of the book is practical but not merely procedural. Each topic is explained as a concept that supports later laboratory work. "
        "Students should read with a notebook nearby, because the best learning happens when explanation, wiring, command use, and observation are connected.",
    )
    doc.add_page_break()

    p(doc, "Textbook Structure", "Heading 1")
    p(
        doc,
        "Each unit is divided into subtopics. The subtopics are written as short textbook sections rather than worksheet blocks. "
        "Where comparison is useful, tables are included. Where a practical caution matters, it is written naturally inside the explanation instead of being separated as an activity or review question.",
    )
    table(
        doc,
        [
            ["Learning Area", "Why It Matters"],
            ["Configuration", "A correctly configured Raspberry Pi is the foundation for remote work and programming."],
            ["Breadboard circuits", "Most beginner experiments are built temporarily before they are made permanent."],
            ["Sensors", "IoT systems depend on trustworthy physical measurements."],
            ["Electronics", "Circuit behavior must be understood before software can be blamed."],
            ["GPIO programming", "Python connects the logic of a program to physical devices."],
        ],
    )
    doc.add_page_break()


def write_unit(doc, unit_no, title, sections):
    p(doc, f"{unit_no}: {title}", "Heading 1")
    p(
        doc,
        f"{title} forms one major part of the course. The sections in this unit are ordered so that a learner can move from the basic idea to practical understanding. "
        "The emphasis is on clear explanation, correct terminology, and careful laboratory thinking.",
    )
    if unit_no == "Unit 01":
        table(
            doc,
            [
                ["Remote Method", "Interface", "Main Use"],
                ["SSH", "Command line", "Run commands, install packages, execute scripts"],
                ["VNC", "Graphical desktop", "Use desktop tools and transfer files visually"],
                ["Local setup", "Keyboard, mouse, monitor", "Initial setup or direct troubleshooting"],
            ],
        )
    elif unit_no == "Unit 02":
        table(
            doc,
            [
                ["Breadboard Area", "Connection Pattern", "Typical Use"],
                ["Terminal strip", "Short connected rows", "Component and signal wiring"],
                ["Power rail", "Long supply path", "Voltage distribution"],
                ["Ground rail", "Common reference path", "Circuit return and signal reference"],
            ],
        )
    elif unit_no == "Unit 03":
        table(
            doc,
            [
                ["Sensor Term", "Meaning"],
                ["Range", "Minimum to maximum measurable value"],
                ["Sensitivity", "Output change for a given input change"],
                ["Accuracy", "Closeness to the true value"],
                ["Response time", "Delay before the output follows the input"],
            ],
        )
    elif unit_no == "Unit 04":
        table(
            doc,
            [
                ["Quantity", "Meaning", "Common Unit"],
                ["Voltage", "Potential difference", "Volt"],
                ["Current", "Charge flow", "Ampere"],
                ["Resistance", "Opposition to current", "Ohm"],
                ["Frequency", "Cycles per second", "Hertz"],
            ],
        )
    elif unit_no == "Unit 05":
        table(
            doc,
            [
                ["GPIO Zero Device", "Hardware Example"],
                ["LED", "Digital output"],
                ["Button", "Digital input"],
                ["MCP3008", "Analog-to-digital conversion"],
                ["MotionSensor", "PIR motion input"],
                ["DistanceSensor", "Ultrasonic distance input"],
            ],
        )
    doc.add_page_break()

    for heading, body in sections:
        p(doc, f"{unit_no} - {heading}", "Heading 1")
        p(doc, body)
        p(doc, "Textbook Discussion", "Heading 2")
        p(
            doc,
            "The topic should be understood as part of a complete IoT workflow. In the early stage of learning, students often separate configuration, wiring, electronics, and programming as if they are unrelated. In practice, they support one another. A remote login depends on correct network configuration. A sensor reading depends on both software and wiring. A Python program is only reliable when the circuit connected to it is electrically sensible.",
        )
        p(doc, "Application in the Laboratory", "Heading 2")
        p(
            doc,
            "When this topic appears in laboratory work, the student should describe the expected result before testing. For a configuration task, the result may be a successful connection or visible desktop. For a circuit task, it may be a measured voltage, glowing LED, stable input, or readable sensor value. For a programming task, it may be a Python output that matches the physical state of the circuit.",
        )
        p(doc, "Common Interpretation", "Heading 2")
        p(
            doc,
            "A frequent beginner mistake is to treat a failure as mysterious. Most failures can be divided into a few categories: no power, wrong ground, wrong pin, wrong command, missing software package, incorrect component orientation, or misunderstanding of the signal. A textbook approach teaches the student to inspect these causes in order instead of randomly changing the circuit or code.",
            "Textbook Note",
        )
        doc.add_page_break()


def appendices(doc):
    p(doc, "Appendix A: Raspberry Pi Configuration Checklist", "Heading 1")
    bullets(
        doc,
        [
            "Raspberry Pi OS has been written correctly to the microSD card.",
            "Network connection is available.",
            "SSH is enabled when terminal access is required.",
            "VNC is enabled when desktop access is required.",
            "The current IP address has been checked.",
            "The username and password have been tested.",
        ],
    )
    doc.add_page_break()

    p(doc, "Appendix B: Breadboard and Sensor Wiring Checklist", "Heading 1")
    bullets(
        doc,
        [
            "Power and ground rails are correctly connected.",
            "Component polarity has been checked.",
            "Sensor pin order has been verified from the module or datasheet.",
            "Pull-up or pull-down resistors are used where required.",
            "GPIO voltage levels are safe for Raspberry Pi.",
            "The circuit has been inspected before power is applied.",
        ],
    )
    doc.add_page_break()

    p(doc, "Appendix C: Formula and Term Reference", "Heading 1")
    table(
        doc,
        [
            ["Formula or Term", "Meaning"],
            ["V = I x R", "Ohm's law relating voltage, current, and resistance"],
            ["P = V x I", "Electrical power"],
            ["Hz", "Cycles per second"],
            ["t0.9", "Approximate time to reach 90 percent of final response"],
            ["FSO", "Full-scale output"],
        ],
    )
    doc.add_page_break()

    p(doc, "Appendix D: Laboratory Report Format", "Heading 1")
    p(
        doc,
        "A laboratory report should record the purpose of the experiment, the components used, the circuit diagram, the commands or program, the observations, and the conclusion. It should also record mistakes because corrected mistakes are part of engineering learning.",
    )
    bullets(
        doc,
        [
            "Title and objective",
            "Components and tools",
            "Circuit diagram",
            "Procedure",
            "Program or command sequence",
            "Observation and result",
            "Errors found and corrected",
            "Conclusion",
        ],
    )


def enrichment_sections(doc):
    sections = [
        (
            "Textbook Extension 1: From Computer to IoT Node",
            "A Raspberry Pi becomes an IoT node when it is configured to sense, process, communicate, and respond. The board alone is only a small computer. The IoT behavior appears when external devices are connected and the program begins to use physical information. This is why configuration and circuit learning belong in the same textbook. Remote access prepares the board for development, while GPIO and sensors prepare it for interaction with the surrounding environment.",
        ),
        (
            "Textbook Extension 2: Building a Reliable Setup Routine",
            "A reliable setup routine reduces confusion in the laboratory. The student should power the board, confirm booting, check network connection, verify the IP address, test SSH or VNC, and then open the required project folder. When this order becomes familiar, later experiments begin smoothly. Many technical problems are avoided simply because the learner follows the same preparation pattern every time.",
        ),
        (
            "Textbook Extension 3: Understanding Remote Workflows",
            "Remote workflows are common in IoT because devices may not always have a monitor, keyboard, or mouse attached. SSH is efficient for commands and scripts, while VNC is convenient for graphical interaction. A student should be comfortable with both. The important idea is that remote access is not separate from IoT; it is one of the practical methods by which IoT devices are maintained.",
        ),
        (
            "Textbook Extension 4: Breadboard Thinking",
            "A breadboard should be read like a map. The learner must know which holes are internally connected and which holes are isolated. A neat circuit layout is easier to debug because the signal path can be followed visually. Breadboard thinking is the habit of seeing both the physical component placement and the hidden electrical connections at the same time.",
        ),
        (
            "Textbook Extension 5: Sensor Wiring Discipline",
            "Sensor circuits depend on correct power, ground, and signal connections. A sensor may appear defective when the real problem is reversed wiring, missing common ground, or a floating data line. DHT sensors are useful for beginners because they show the importance of VCC, GND, DATA, and pull-up behavior in a compact example.",
        ),
        (
            "Textbook Extension 6: Reading Sensor Data Carefully",
            "A sensor value should not be accepted blindly. The student should ask whether the value is inside the expected range, whether it changes when the physical condition changes, and whether the reading is stable. Sensor fundamentals such as range, sensitivity, accuracy, hysteresis, and response time are practical tools for judging whether a measurement is believable.",
        ),
        (
            "Textbook Extension 7: Calibration as a Habit",
            "Calibration is not only a procedure for advanced instruments. Even beginner projects benefit from comparing a sensor reading with a known condition. A temperature reading can be compared with a trusted thermometer, and a distance reading can be compared with a ruler. Calibration teaches that measured data has context and uncertainty.",
        ),
        (
            "Textbook Extension 8: Circuit Quantities in Practice",
            "Voltage, current, and resistance are not abstract textbook words. They explain why an LED needs a resistor, why a GPIO pin has current limits, why a missing ground breaks communication, and why supply voltage must match the device. Ohm's law gives the learner a way to predict whether a simple circuit is reasonable before it is powered.",
        ),
        (
            "Textbook Extension 9: Components as Building Blocks",
            "Resistors, capacitors, diodes, and transistors are building blocks. A resistor controls current or defines a signal state. A capacitor stores charge and smooths changes. A diode controls current direction or protects against reverse voltage. A transistor allows a small signal to control a larger load. Understanding these roles makes schematic diagrams less intimidating.",
        ),
        (
            "Textbook Extension 10: From Analog Circuits to Digital Logic",
            "Operational amplifiers and digital ICs show two sides of electronics. Analog circuits shape and compare continuous signals. Digital circuits store, select, buffer, and transfer logic values. IoT systems often combine both: sensors may begin with analog behavior, while the computer processes digital information.",
        ),
        (
            "Textbook Extension 11: Python as Hardware Language",
            "In GPIO Zero, Python becomes a language for physical behavior. An LED object represents an output device, a Button object represents an input device, and sensor objects represent measured conditions. This object-based style helps beginners connect code statements with real components on the breadboard.",
        ),
        (
            "Textbook Extension 12: The Complete Beginner IoT Path",
            "The complete path of this course can be summarized as prepare, connect, measure, understand, and program. Prepare the Raspberry Pi. Connect circuits on the breadboard. Measure the physical world with sensors. Understand the electronic principles behind the circuit. Program the hardware using GPIO Zero. These steps form the foundation for more advanced IoT study.",
        ),
    ]
    for title, body in sections:
        p(doc, title, "Heading 1")
        p(doc, body)
        p(
            doc,
            "This section strengthens the textbook flow by connecting earlier course topics into a broader explanation. It is written for reading and understanding, not as a worksheet or activity block.",
            "Textbook Note",
        )
        doc.add_page_break()


def main():
    doc = Document()
    configure(doc)
    front_matter(doc)
    for unit_no, title, sections in UNITS:
        write_unit(doc, unit_no, title, sections)
    enrichment_sections(doc)
    appendices(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
