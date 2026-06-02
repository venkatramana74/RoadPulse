from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name(
    "Basic_Circuits_and_Raspberry_Pi_Configuration_Textbook_100_Pages.docx"
)

UNITS = [
    (
        "UNIT 01: Raspberry Pi Configuration",
        [
            "Raspberry Pi in IoT",
            "Raspberry Pi Board Overview",
            "GPIO Header and Pin Groups",
            "Preparing the SD Card",
            "Installing Raspberry Pi OS",
            "First Boot and Basic Setup",
            "Raspberry Pi Configuration Tool",
            "Enabling SSH",
            "Finding the IP Address",
            "Using PuTTY for SSH",
            "Enabling VNC",
            "Using VNC Viewer",
            "File Transfer with VNC",
            "Linux Terminal Basics",
            "Package Update and Software Installation",
            "Remote Access Troubleshooting",
            "Raspberry Pi Safety Practices",
            "Unit 01 Review and Practice",
        ],
    ),
    (
        "UNIT 02: Breadboard and Basic Sensor Circuits",
        [
            "Introduction to Breadboards",
            "Breadboard Internal Connections",
            "Power Rails and Ground Rails",
            "Component Placement",
            "Jumper Wires",
            "Why Breadboards Help Learning",
            "Resistor Use on Breadboards",
            "LED Circuit Basics",
            "Pull-Up Resistors",
            "Pull-Down Resistors",
            "DHT Sensor Family",
            "DHT Pin Functions",
            "DHT Pull-Up Requirement",
            "Building a DHT Circuit",
            "Breadboard Debugging",
            "Clean Wiring Practice",
            "Unit 02 Review and Practice",
        ],
    ),
    (
        "UNIT 03: Sensor Fundamentals",
        [
            "What Is a Sensor?",
            "Sensor Input and Output",
            "Analog Sensors",
            "Digital Sensors",
            "Sensor Range",
            "Sensitivity",
            "Accuracy and Error",
            "Linearity",
            "Hysteresis",
            "Response Time",
            "Calibration",
            "Noise and Stability",
            "Sensor Selection",
            "Temperature Sensors",
            "Humidity Sensors",
            "Motion and Distance Sensors",
            "Unit 03 Review and Practice",
        ],
    ),
    (
        "UNIT 04: Basic Electronics and Circuits",
        [
            "Electricity by Water Analogy",
            "Voltage",
            "Current",
            "Resistance",
            "Ohm's Law",
            "DC Signals",
            "AC Signals",
            "Frequency and Hertz",
            "Series Circuits",
            "Parallel Circuits",
            "Dry Cells and Supply Voltage",
            "Capacitors",
            "Smoothing Circuits",
            "Timer Circuits",
            "Low-Pass Filters",
            "High-Pass Filters",
            "Diode Basics",
            "Types of Diodes",
            "Zener Diodes",
            "Transistor Basics",
            "NPN and PNP Transistors",
            "FET and MOSFET Overview",
            "Back EMF and Protection Diodes",
            "Circuit Symbols",
            "Wire Crossings and Junctions",
            "Unit Abbreviations in Diagrams",
            "Operational Amplifiers",
            "Inverting and Non-Inverting Circuits",
            "Comparators",
            "Digital ICs",
            "Flip-Flops",
            "Address and Data Buses",
            "Decoders and Buffers",
            "Unit 04 Review and Practice",
        ],
    ),
    (
        "UNIT 05: GPIO Zero Programming",
        [
            "Introduction to GPIO Zero",
            "Installing GPIO Zero",
            "Blinking One LED",
            "LED Program with Functions",
            "Blinking Three LEDs",
            "MCP3008 and Potentiometer",
            "Button Input",
            "Button Game",
            "Motion Sensor Program",
            "Distance Sensor Program",
            "Good Python Hardware Style",
            "Unit 05 Review and Practice",
        ],
    ),
]

EXPLANATIONS = {
    "Raspberry Pi": "Raspberry Pi is a compact single-board computer that can run Linux, connect to networks, and control external electronic circuits through its GPIO header.",
    "GPIO": "GPIO pins are general-purpose input and output pins. They allow the board to read switches and sensors or control LEDs, modules, and other digital devices.",
    "SSH": "SSH provides secure command-line access from another computer. It is useful when the Raspberry Pi is used without a separate monitor or keyboard.",
    "VNC": "VNC provides remote desktop access. It allows the Raspberry Pi graphical desktop to be viewed and controlled from another computer.",
    "Breadboard": "A breadboard is a solderless prototyping board. It helps learners build, test, modify, and reuse circuits without permanently soldering components.",
    "DHT": "DHT11 and DHT22 sensors measure temperature and humidity. Their circuits usually include VCC, GND, DATA, and sometimes an unused pin.",
    "Sensor": "A sensor converts a physical quantity such as temperature, humidity, motion, distance, pressure, or light into an electrical or digital signal.",
    "Hysteresis": "Hysteresis describes the difference in sensor output when the measured value approaches the same point from different directions.",
    "Response": "Response time is the delay between a change in physical input and the sensor output reaching its expected value.",
    "Voltage": "Voltage is electric potential difference. It acts like the push that can drive current through a closed circuit.",
    "Current": "Current is the flow of electric charge. It must be controlled so that components and Raspberry Pi GPIO pins are not damaged.",
    "Resistance": "Resistance opposes current flow. Resistors are used for current limiting, voltage division, and pull-up or pull-down circuits.",
    "Ohm": "Ohm's law, V = I x R, connects voltage, current, and resistance and is the foundation for many beginner circuit calculations.",
    "Diode": "A diode conducts more easily in one direction than the other. Its terminals are called anode and cathode.",
    "Transistor": "A transistor is used as an electronic switch or amplifier. It allows a small control signal to influence a larger current path.",
    "Op": "An operational amplifier has differential inputs and an output. It is widely used in analog signal conditioning and comparison circuits.",
    "GPIO Zero": "GPIO Zero is a beginner-friendly Python library for Raspberry Pi physical computing. It provides simple classes for LEDs, buttons, sensors, and other devices.",
    "Button": "A button is a digital input device. Correct pull-up or pull-down behavior keeps the input stable when the button is not pressed.",
    "Distance": "A distance sensor estimates how far an object is from the sensor. Ultrasonic sensors do this by timing sound pulses.",
}

COMBINE_WITH_NEXT = {
    "Finding the IP Address",
    "Enabling VNC",
    "Using VNC Viewer",
    "Package Update and Software Installation",
    "Raspberry Pi Safety Practices",
    "Voltage",
    "Resistance",
    "DC Signals",
    "Series Circuits",
    "Low-Pass Filters",
    "Diode Basics",
    "Transistor Basics",
    "FET and MOSFET Overview",
    "Operational Amplifiers",
    "Comparators",
    "Button Input",
}


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04

    for style_name, size, color, before, after in [
        ("Title", 24, "0B2545", 0, 8),
        ("Heading 1", 16, "1F4D78", 10, 5),
        ("Heading 2", 13, "2E74B5", 7, 3),
        ("Heading 3", 11.5, "1F4D78", 5, 2),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Page Title" not in styles:
        style = styles.add_style("Page Title", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Heading 1"]
        style.font.size = Pt(15)
        style.font.color.rgb = RGBColor.from_string("0B2545")
        style.font.bold = True
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)

    if "Small Note" not in styles:
        style = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor.from_string("555555")
        style.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.text = "Basic Circuits and Raspberry Pi Configuration for IoT Applications"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string("666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    run.font.size = Pt(8.5)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def add_p(doc, text, style="Normal"):
    para = doc.add_paragraph(style=style)
    para.add_run(text)
    return para


def add_bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(2)
        para.add_run(item)


def add_numbers(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(2)
        para.add_run(item)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for col, text in enumerate(rows[0]):
        cell = table.rows[0].cells[col]
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        shade(cell, "F2F4F7")
    for row in rows[1:]:
        cells = table.add_row().cells
        for col, text in enumerate(row):
            cells[col].text = text
            cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_p(doc, "", "Small Note")


def explanation_for(topic):
    for key, value in EXPLANATIONS.items():
        if key.lower() in topic.lower():
            return value
    return (
        f"{topic} is an important part of basic IoT hardware learning. "
        "It should be studied as a practical skill, not only as a definition."
    )


def lesson_page(doc, unit, topic, index):
    add_p(doc, unit, "Small Note")
    add_p(doc, topic, "Page Title")
    add_p(doc, "Learning Objectives", "Heading 3")
    add_bullets(
        doc,
        [
            f"Explain the meaning of {topic.lower()} in an IoT laboratory context.",
            "Connect the topic to Raspberry Pi configuration, circuit practice, or sensor use.",
            "Apply the idea safely during a small practical experiment.",
        ],
    )
    add_p(doc, "Concept Explanation", "Heading 2")
    add_p(doc, explanation_for(topic))
    add_p(
        doc,
        "A beginner should learn this topic by connecting the theory to a visible action. "
        "For example, a command should produce a clear terminal result, a wire should connect two known points, "
        "and a sensor reading should be interpreted using its electrical behavior. This habit prevents blind trial and error.",
    )
    add_p(doc, "Practical Notes", "Heading 2")
    add_bullets(
        doc,
        [
            "Record pin numbers, component values, commands, and observations in a lab notebook.",
            "Check power and ground before assuming that the program is wrong.",
            "Test one change at a time so the cause of success or failure is visible.",
            "Use correct voltage levels when connecting anything to Raspberry Pi GPIO.",
        ],
    )
    if index % 7 == 0:
        add_p(doc, "Reference Table", "Heading 3")
        add_table(
            doc,
            [
                ["Item", "Meaning", "Practical Use"],
                ["Input", "Signal read by the board or circuit", "Buttons, sensors, and data lines"],
                ["Output", "Signal produced by the board or circuit", "LEDs, control pins, and modules"],
                ["Ground", "Common reference point", "Stable and safe circuit operation"],
            ],
        )
    add_p(doc, "Practical Activity", "Heading 3")
    add_numbers(
        doc,
        [
            "Write the purpose of the topic in one sentence.",
            "Draw the smallest circuit, command sequence, or block diagram needed to test it.",
            "Predict one correct result and one likely beginner error.",
        ],
    )
    add_p(doc, "Review Questions", "Heading 3")
    add_numbers(
        doc,
        [
            f"Why is {topic.lower()} important in this course?",
            "What mistake could damage the circuit or prevent the experiment from working?",
            "How would you verify that the topic has been understood in practice?",
        ],
    )


def combined_lesson_page(doc, unit, first_topic, second_topic, index):
    topic = f"{first_topic} and {second_topic}"
    add_p(doc, unit, "Small Note")
    add_p(doc, topic, "Page Title")
    add_p(doc, "Learning Objectives", "Heading 3")
    add_bullets(
        doc,
        [
            f"Explain the connection between {first_topic.lower()} and {second_topic.lower()}.",
            "Apply both ideas during configuration, wiring, measurement, or programming practice.",
            "Recognize the common mistakes that appear when the two topics are confused.",
        ],
    )
    add_p(doc, first_topic, "Heading 2")
    add_p(doc, explanation_for(first_topic))
    add_p(doc, second_topic, "Heading 2")
    add_p(doc, explanation_for(second_topic))
    add_p(doc, "Practical Connection", "Heading 2")
    add_p(
        doc,
        "These two subtopics are studied together because they appear naturally in the same laboratory workflow. "
        "A learner should not memorize them as isolated definitions. Instead, the learner should ask how the first "
        "idea supports the second, what must be checked before testing, and what observation proves that the setup is correct.",
    )
    if index % 5 == 0:
        add_p(doc, "Reference Table", "Heading 3")
        add_table(
            doc,
            [
                ["Check", "Reason", "Evidence"],
                ["Connection", "The circuit or network path must be complete", "Device responds or signal is visible"],
                ["Value", "Voltage, pin, address, or command must be correct", "Measured or displayed result matches expectation"],
                ["Safety", "Incorrect wiring can damage parts", "Power and ground are verified before testing"],
            ],
        )
    add_p(doc, "Practical Activity", "Heading 3")
    add_numbers(
        doc,
        [
            f"Write one sentence for {first_topic.lower()} and one for {second_topic.lower()}.",
            "Draw a small workflow showing how the two ideas are used together.",
            "List one likely mistake and one method to detect it.",
        ],
    )
    add_p(doc, "Review Questions", "Heading 3")
    add_numbers(
        doc,
        [
            f"How does {first_topic.lower()} support {second_topic.lower()}?",
            "What practical check should be completed before testing?",
            "What observation would show that both subtopics are understood?",
        ],
    )


def front_matter(doc):
    add_p(doc, "Basic Circuits and Raspberry Pi Configuration for IoT Applications", "Title")
    add_p(
        doc,
        "A practical textbook based on SIC IoT Chapter 2: Basic Circuit and Raspberry Pi Configuration",
        "Heading 2",
    )
    add_p(
        doc,
        "This manuscript expands the course topics into original textbook language. "
        "It follows the structure of Raspberry Pi configuration, breadboard practice, sensor fundamentals, "
        "basic electronics, and GPIO Zero programming. It is intended for publication preparation, classroom use, "
        "and student laboratory study.",
    )
    add_table(
        doc,
        [
            ["Book Element", "Description"],
            ["Audience", "Beginners in IoT, electronics, Raspberry Pi, and physical computing"],
            ["Approach", "Theory, practical procedure, safety notes, review questions, and lab tasks"],
            ["Course Flow", "Unit 01 to Unit 05, from configuration to programming"],
        ],
    )
    doc.add_page_break()

    add_p(doc, "Preface", "Page Title")
    add_p(
        doc,
        "The Internet of Things begins with simple but disciplined skills: configuring a board, "
        "understanding a circuit, reading a sensor, and writing a program that responds to the physical world. "
        "This textbook introduces those skills through Raspberry Pi configuration and basic circuit practice.",
    )
    add_p(
        doc,
        "The explanations are written for beginners. Each page is a teachable unit with objectives, notes, "
        "activities, and review questions. Instructors can expand each page into a lecture or lab session, "
        "while students can use the pages as a guided study manual.",
    )
    doc.add_page_break()

    add_p(doc, "How to Use This Textbook", "Page Title")
    add_bullets(
        doc,
        [
            "Study each unit in order during the first reading.",
            "Complete practical work only after checking power, ground, polarity, and pin numbers.",
            "Maintain a lab notebook with commands, circuit diagrams, observations, and corrections.",
            "Use review questions for self-assessment before moving to the next lesson.",
            "Treat Raspberry Pi GPIO pins carefully; wrong voltage or short circuits can damage the board.",
        ],
    )
    doc.add_page_break()

    add_p(doc, "Table of Contents", "Page Title")
    add_table(
        doc,
        [
            ["Part", "Topic", "Role in the Book"],
            ["Unit 01", "Raspberry Pi Configuration", "Prepare the board and enable remote access"],
            ["Unit 02", "Breadboard and Sensor Circuits", "Build safe temporary circuits"],
            ["Unit 03", "Sensor Fundamentals", "Understand sensor behavior and characteristics"],
            ["Unit 04", "Basic Electronics and Circuits", "Learn components and circuit reading"],
            ["Unit 05", "GPIO Zero Programming", "Control hardware with Python"],
            ["Appendices", "Checklists and Review", "Support laboratory and examination preparation"],
        ],
    )
    doc.add_page_break()

    add_p(doc, "Course Map", "Page Title")
    add_p(
        doc,
        "The course moves from configuration to construction, then from construction to measurement and programming. "
        "This order is important because practical IoT work requires both software confidence and circuit discipline.",
    )
    add_table(
        doc,
        [
            ["Stage", "Main Skill", "Outcome"],
            ["Configuration", "Install and access Raspberry Pi", "Ready development environment"],
            ["Prototyping", "Use breadboard and components", "Safe circuit construction"],
            ["Measurement", "Understand sensors", "Reliable data interpretation"],
            ["Electronics", "Read and analyze circuits", "Better troubleshooting"],
            ["Programming", "Use GPIO Zero", "Interactive physical computing"],
        ],
    )
    doc.add_page_break()


def unit_intro(doc, unit):
    add_p(doc, unit, "Page Title")
    add_p(
        doc,
        "This unit introduces a major block of the course. Read the pages in sequence, perform the practical tasks, "
        "and answer the review questions. The aim is to develop habits that are useful in real IoT laboratories: "
        "careful wiring, correct configuration, systematic testing, and clear documentation.",
    )
    add_bullets(
        doc,
        [
            "Start with the concept before touching the hardware.",
            "Verify each connection or command before moving ahead.",
            "Write down what worked, what failed, and how it was corrected.",
        ],
    )


def appendices(doc):
    items = [
        (
            "Appendix A: Raspberry Pi Configuration Checklist",
            [
                "Raspberry Pi OS installed",
                "Network connected",
                "SSH enabled when terminal access is needed",
                "VNC enabled when desktop access is needed",
                "IP address recorded",
                "Username and password tested",
                "Package list updated",
            ],
        ),
        (
            "Appendix B: Breadboard Wiring Checklist",
            [
                "Power rail connected correctly",
                "Ground rail connected correctly",
                "No unintended short circuit",
                "LED polarity checked",
                "Resistor value checked",
                "Sensor pinout checked",
                "GPIO voltage confirmed safe",
            ],
        ),
        (
            "Appendix C: Sensor Evaluation Sheet",
            [
                "Measured quantity",
                "Range",
                "Accuracy",
                "Sensitivity",
                "Response time",
                "Output type",
                "Supply voltage",
                "Interface",
                "Calibration need",
            ],
        ),
        (
            "Appendix D: Formula Reference",
            [
                "Ohm's law: V = I x R",
                "Power: P = V x I",
                "Series resistance: Rtotal = R1 + R2 + ...",
                "Parallel resistance: 1/Rtotal = 1/R1 + 1/R2 + ...",
                "Frequency: hertz means cycles per second",
            ],
        ),
        (
            "Appendix E: GPIO Zero Quick Reference",
            [
                "Import the required class",
                "Assign the correct GPIO pin",
                "Use on, off, blink, or value reading methods",
                "Keep the program running when waiting for input",
                "Stop the program safely after testing",
            ],
        ),
        (
            "Appendix F: Laboratory Report Format",
            [
                "Title",
                "Objective",
                "Components used",
                "Circuit diagram",
                "Procedure",
                "Program or commands",
                "Observation",
                "Result",
                "Errors corrected",
                "Conclusion",
            ],
        ),
        (
            "Appendix G: Final Review Questions",
            [
                "What is the difference between SSH and VNC?",
                "Why is a breadboard useful?",
                "Why do some sensors need pull-up resistors?",
                "What is hysteresis?",
                "How are DC and AC different?",
                "Why should LEDs use resistors?",
                "What is the role of an ADC?",
                "How does GPIO Zero simplify programming?",
            ],
        ),
    ]
    for index, (title, points) in enumerate(items):
        add_p(doc, title, "Page Title")
        add_p(doc, "Use this page as a practical handout, revision page, or instructor checklist.")
        add_bullets(doc, points)
        add_p(doc, "Instructor Note", "Heading 2")
        add_p(
            doc,
            "Students should complete this page with observations from their own Raspberry Pi and circuit setup. "
            "The checklist is most useful when it is filled during real laboratory work rather than after the experiment.",
        )
        if index != len(items) - 1:
            doc.add_page_break()


def main():
    doc = Document()
    configure(doc)
    front_matter(doc)
    page_no = 5
    lesson_index = 0
    for unit, topics in UNITS:
        unit_intro(doc, unit)
        doc.add_page_break()
        page_no += 1
        i = 0
        while i < len(topics):
            topic = topics[i]
            lesson_index += 1
            if topic in COMBINE_WITH_NEXT and i + 1 < len(topics):
                combined_lesson_page(doc, unit, topic, topics[i + 1], lesson_index)
                i += 2
            else:
                lesson_page(doc, unit, topic, lesson_index)
                i += 1
            doc.add_page_break()
            page_no += 1
    appendices(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
